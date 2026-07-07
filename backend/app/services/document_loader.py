import os
from pathlib import Path
from typing import List
from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from docx import Document as DocxDocument

class DocumentLoader:
    SUPPORTED = {".pdf", ".txt", ".docx"}
    
    def load(self, file_path: str) -> List[Document]:
        ext = Path(file_path).suffix.lower()
        if ext not in self.SUPPORTED:
            raise ValueError(f"Unsupported file type: {ext}")
        
        if ext == ".pdf":
            return PyPDFLoader(file_path).load()
        elif ext == ".txt":
            return TextLoader(file_path, encoding="utf-8").load()
        elif ext == ".docx":
            return self._load_docx(file_path)
    
    def _load_docx(self, path: str) -> List[Document]:
        doc = DocxDocument(path)
        text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        return [Document(page_content=text, metadata={"source": path})]